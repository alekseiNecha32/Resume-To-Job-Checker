"""
Resume converter service for generating styled DOCX files using Pandoc.
"""
import os
import json
import tempfile
from pathlib import Path
from jinja2 import Template
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

try:
    import pypandoc
except ImportError:
    pypandoc = None


TEMPLATES_DIR = Path(__file__).parent.parent / "templates" / "resume"


class PandocNotAvailableError(Exception):
    """Raised when Pandoc is not installed."""
    pass


class TemplateNotFoundError(Exception):
    """Raised when a template is not found."""
    pass


def get_available_templates() -> list:
    """
    Get list of available resume templates.

    Returns:
        List of template metadata dictionaries.
    """
    templates = []

    if not TEMPLATES_DIR.exists():
        return templates

    for template_dir in TEMPLATES_DIR.iterdir():
        if template_dir.is_dir():
            metadata_file = template_dir / "metadata.json"
            if metadata_file.exists():
                with open(metadata_file, "r", encoding="utf-8") as f:
                    metadata = json.load(f)
                    templates.append(metadata)

    return templates


def get_template_path(template_id: str) -> Path:
    """Get the path to a template directory."""
    template_path = TEMPLATES_DIR / template_id
    if not template_path.exists():
        raise TemplateNotFoundError(f"Template '{template_id}' not found")
    return template_path


def resume_to_markdown(resume_data: dict, template_id: str) -> str:
    """
    Convert resume JSON to Markdown using Jinja2 template.

    Args:
        resume_data: The resume dictionary from frontend
        template_id: The template ID to use

    Returns:
        Rendered markdown string
    """
    template_path = get_template_path(template_id)
    template_file = template_path / "template.md"

    if not template_file.exists():
        raise TemplateNotFoundError(f"Template file not found: {template_file}")

    with open(template_file, "r", encoding="utf-8") as f:
        template_str = f.read()

    template = Template(template_str)

    # Normalize data
    context = {
        "name": resume_data.get("name", "Your Name"),
        "title": resume_data.get("title", ""),
        "contact": resume_data.get("contact", []),
        "sections": resume_data.get("sections", [])
    }

    return template.render(**context)


def _create_reference_docx(template_id: str, output_path: Path) -> None:
    """
    Create a reference DOCX file with proper styling for a template.

    Args:
        template_id: The template ID
        output_path: Path to save the reference document
    """
    doc = Document()

    # Template-specific styling
    if template_id == "classic":
        font_name = "Times New Roman"
        body_size = 11
        heading1_size = 18
        heading2_size = 14
        margins = 1.0  # inches
    elif template_id == "modern":
        font_name = "Calibri"
        body_size = 11
        heading1_size = 16
        heading2_size = 12
        margins = 0.75
    else:  # compact
        font_name = "Arial"
        body_size = 10
        heading1_size = 14
        heading2_size = 11
        margins = 0.5

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(margins)
        section.bottom_margin = Inches(margins)
        section.left_margin = Inches(margins)
        section.right_margin = Inches(margins)

    # Configure Normal style
    normal_style = doc.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(body_size)

    # Configure Heading 1
    h1_style = doc.styles["Heading 1"]
    h1_style.font.name = font_name
    h1_style.font.size = Pt(heading1_size)
    h1_style.font.bold = True
    if template_id == "classic":
        h1_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Configure Heading 2
    h2_style = doc.styles["Heading 2"]
    h2_style.font.name = font_name
    h2_style.font.size = Pt(heading2_size)
    h2_style.font.bold = True
    if template_id == "compact":
        h2_style.font.all_caps = True

    # Configure List Bullet style
    try:
        list_style = doc.styles["List Bullet"]
        list_style.font.name = font_name
        list_style.font.size = Pt(body_size)
    except KeyError:
        pass

    doc.save(output_path)


def _ensure_reference_docx(template_id: str) -> Path:
    """
    Ensure the reference DOCX exists for a template, creating it if needed.

    Args:
        template_id: The template ID

    Returns:
        Path to the reference DOCX file
    """
    template_path = get_template_path(template_id)
    reference_path = template_path / "reference.docx"

    if not reference_path.exists():
        _create_reference_docx(template_id, reference_path)

    return reference_path


def generate_docx(resume_data: dict, template_id: str) -> bytes:
    """
    Generate a styled DOCX file from resume data.

    Args:
        resume_data: The resume dictionary from frontend
        template_id: The template ID to use

    Returns:
        DOCX file as bytes
    """
    if pypandoc is None:
        raise PandocNotAvailableError("pypandoc is not installed")

    # Ensure reference document exists
    reference_path = _ensure_reference_docx(template_id)

    # Convert resume to markdown
    markdown_content = resume_to_markdown(resume_data, template_id)

    # Create temp file for output
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        # Convert markdown to DOCX using Pandoc
        pypandoc.convert_text(
            markdown_content,
            "docx",
            format="md",
            outputfile=tmp_path,
            extra_args=[f"--reference-doc={reference_path}"]
        )

        # Read the generated file
        with open(tmp_path, "rb") as f:
            content = f.read()

        return content
    finally:
        # Clean up temp file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def generate_docx_fallback(resume_data: dict, template_id: str) -> bytes:
    """
    Generate a styled DOCX file using python-docx directly (fallback if Pandoc unavailable).

    Args:
        resume_data: The resume dictionary from frontend
        template_id: The template ID to use

    Returns:
        DOCX file as bytes
    """
    import io

    # Template-specific styling
    if template_id == "classic":
        font_name = "Times New Roman"
        body_size = 11
        heading1_size = 18
        heading2_size = 14
        margins = 1.0
        center_name = True
    elif template_id == "modern":
        font_name = "Calibri"
        body_size = 11
        heading1_size = 16
        heading2_size = 12
        margins = 0.75
        center_name = False
    else:  # compact
        font_name = "Arial"
        body_size = 10
        heading1_size = 14
        heading2_size = 11
        margins = 0.5
        center_name = False

    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(margins)
        section.bottom_margin = Inches(margins)
        section.left_margin = Inches(margins)
        section.right_margin = Inches(margins)

    # Configure Normal style
    normal_style = doc.styles["Normal"]
    normal_style.font.name = font_name
    normal_style.font.size = Pt(body_size)

    # Add name
    name = resume_data.get("name", "Your Name")
    p_name = doc.add_paragraph()
    if center_name:
        p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_name = p_name.add_run(name)
    run_name.bold = True
    run_name.font.size = Pt(heading1_size)
    run_name.font.name = font_name

    # Add contact info
    for line in resume_data.get("contact", []):
        p = doc.add_paragraph()
        if center_name:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(body_size)

    doc.add_paragraph("")

    # Add sections
    for section in resume_data.get("sections", []):
        title = (section.get("title") or "").strip()
        items = section.get("items") or []
        if not title or not items:
            continue

        # Section heading
        h = doc.add_paragraph()
        run = h.add_run(title if template_id != "compact" else title.upper())
        run.bold = True
        run.font.size = Pt(heading2_size)
        run.font.name = font_name

        # Items
        if section.get("id") == "professional-summary":
            text = " ".join(i.get("text", "").strip() for i in items if i.get("text"))
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = font_name
                run.font.size = Pt(body_size)
        else:
            for it in items:
                text = (it.get("text") or "").strip()
                if not text:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(text)
                run.font.name = font_name
                run.font.size = Pt(body_size)

        doc.add_paragraph("")

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
