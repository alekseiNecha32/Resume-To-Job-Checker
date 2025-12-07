from flask import Blueprint, request, send_file, jsonify
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io

export_bp = Blueprint("export", __name__, url_prefix="/api/export")


@export_bp.route("/resume-docx", methods=["POST"])
def export_resume_docx():
    data = request.get_json(silent=True) or {}
    resume = data.get("resume")

    if not resume:
        return jsonify({"error": "missing_resume"}), 400

    doc = Document()

    # ----- Base font -----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ===== HEADER =====
    name = resume.get("name") or "Your Name"
    p_name = doc.add_paragraph()
    p_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r_name = p_name.add_run(name)
    r_name.bold = True
    r_name.font.size = Pt(16)

    for line in resume.get("contact", []):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run(line)

    doc.add_paragraph("")  # spacer

    # ===== SECTIONS =====
    for section in resume.get("sections", []):
        title = (section.get("title") or "").strip()
        items = section.get("items") or []
        if not title or not items:
            continue

        # Section heading
        doc.add_paragraph(title, style="Heading 1")

        # Special case: PROFESSIONAL SUMMARY → one paragraph
        if section.get("id") == "professional-summary":
            text = " ".join(i.get("text", "").strip() for i in items if i.get("text"))
            if text:
                doc.add_paragraph(text)
        else:
            # Everything else as bullets (easy to edit later)
            for it in items:
                text = (it.get("text") or "").strip()
                if not text:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(text)

        doc.add_paragraph("")  # spacer

    # ===== Return as .docx =====
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name="optimized_resume.docx",
        mimetype=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )
